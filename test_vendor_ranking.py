import unittest
import pandas as pd
from flask import Flask, request, render_template, jsonify
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import csv
import json
from docx import Document
from docx.shared import Inches

from vendor_ranking import get_relevant_vendors, rank_vendors,extract_pros_text,extract_feature_names

class TestVendorRanking(unittest.TestCase):

    def setUp(self):
        # Load data from CSV file
        df = pd.read_csv("C:/Users/Prahalad M/Downloads/G2 software - CRM Category Product Overviews.csv")
        df1 = df.copy()

        # Replace NaN with empty JSON array in 'Features' column
        df1['Features'] = df1['Features'].fillna('[]')

        # Apply the function to extract feature names
        df1['feature_names'] = df1['Features'].apply(extract_feature_names)

        # Replace NaN with empty JSON array in 'pros_list' column
        df1['pros_list'] = df1['pros_list'].fillna('[]')

        # Apply the function to extract pros text
        df1['pros_text'] = df1['pros_list'].apply(extract_pros_text)

        # Create a new column 'combined_text'
        df1['all_features'] = df1.apply(lambda row: row['pros_text'] if not row['feature_names'] else row['feature_names'], axis=1)

        # Create a new DataFrame with selected columns
        new_df = df1[['product_name', 'seller','main_category', 'categories', 'rating', 'all_features']].copy()
        new_df = new_df.rename(columns={'product_name': 'software_name', 'seller': 'vendor'})
        self.df=new_df.copy()

    def test_get_relevant_vendors_with_matching_keywords(self):
        relevant_vendors = get_relevant_vendors(['workflow', 'automation'], self.df, threshold=0.6)
        # Assert that at least one relevant vendor is found
        self.assertTrue(len(relevant_vendors) > 0)

    def test_get_relevant_vendors_with_no_matching_keywords(self):
        relevant_vendors = get_relevant_vendors(['nonexistent_feature'], self.df, threshold=0.6)
        # Assert that no vendors are found
        self.assertTrue(relevant_vendors.empty)

    def test_rank_vendors_with_rating_weight(self):
        ranked_vendors = rank_vendors(['customer', 'support'], self.df, threshold=0.6, rating_weight=0.2)
        # Assert that the highest-ranked vendor has rank 1
        self.assertEqual(ranked_vendors['rank'].iloc[0], 1)

    def test_rank_vendors_without_rating_weight(self):
        ranked_vendors = rank_vendors(['sales', 'management'], self.df, threshold=0.6, rating_weight=0)
        # Assert that at least one vendor is ranked
        self.assertTrue(len(ranked_vendors) > 0)
    
    def test_rank_vendors_with_no_input(self):
        ranked_vendors = rank_vendors([], self.df, threshold=0.6, rating_weight=0.2)  # Empty keywords list
        #  Assert that the output is not empty
        self.assertFalse(ranked_vendors.empty)

def generate_test_report(test_results):
    #  Print to console
    for test_case, result in test_results.items():
        test_type = "Positive" if test_case.startswith("test_get_relevant_vendors_with_matching") or test_case.startswith("test_rank_vendors_with_rating") else "Negative"
        status = "Pass" if result else "Fail"
        description = test_case.replace('test_', '').replace('_', ' ').capitalize()
        print(f"{test_case}, {test_type}, {status}, {description}")
    total_tests = len(test_results)
    passed_tests = sum(test_results.values())
    failed_tests = total_tests - passed_tests
    print("Total tests:"+str(total_tests)+"\nPassed tests:"+str(passed_tests)+"\nFailed tests:"+str(failed_tests))
    #  Write to Word document
    document = Document()
    document.add_heading('Test Report', 0)  # Add a heading

    #Add overall results
    document.add_paragraph(f"Total tests: {total_tests}")
    document.add_paragraph(f"Passed tests: {passed_tests}")
    document.add_paragraph(f"Failed tests: {failed_tests}")

    # Create table
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Case'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Description'
    
    for test_case, result in test_results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = test_case
        row_cells[1].text = "Positive" if test_case.startswith("test_get_relevant_vendors_with_matching") or test_case.startswith("test_rank_vendors_with_rating") else "Negative"
        row_cells[2].text = "Pass" if result else "Fail"
        row_cells[3].text = test_case.replace('test_', '').replace('_', ' ').capitalize() 
    document.add_paragraph('') 
    document.save('C:/Users/Prahalad M/Downloads/test_report.docx')

    return "Report saved to test_report.docx"
if __name__ == '__main__':
    suite = unittest.TestSuite()
    loader = unittest.TestLoader()
    suite.addTest(loader.loadTestsFromTestCase(TestVendorRanking))
    runner = unittest.TextTestRunner()
    result = runner.run(suite)
    # Get successful test cases
    test_names = [name for name in dir(TestVendorRanking) if name.startswith('test')]
    test_results = {name: False for name in test_names}  # Initialize with False

    if result.wasSuccessful():
        for test_name in test_names:
            test_results[test_name] = True
    else:
        for error in result.errors + result.failures:
            test_name = error[0]._testMethodName
            if test_name in test_results:  # Check if the test exists in test_results
                test_results[test_name] = False
            else:
                print(f"WARNING: Failed test {test_name} not found in initial test list.")

report = generate_test_report(test_results)
